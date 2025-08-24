#!/usr/bin/env python3
"""Generate a minimal DOCX agreement skeleton with required structural elements.
Requires: python-docx (install from requirements.txt)

Usage:
  ./create_sample_agreement.py --out Agreement_v1.docx --campaign TOKEN123

What it does:
  - Creates headings & placeholder body text
  - Inserts a footer with remote logo URL containing campaign token
  - Adds signature block placeholders

Manual follow-up (cannot be automated reliably here):
  - Insert Table of Contents (after opening in Word / LibreOffice)
  - Add cross-references between headings
  - Introduce a deliberate date format inconsistency (edit one date)
  - Optionally embed zero-width watermark using zero_width_watermark.py
"""
import argparse
import datetime
import pathlib
try:
    from docx import Document  # type: ignore
except ImportError as e:  # pragma: no cover
    raise SystemExit('python-docx not installed. Run: pip install python-docx')

def build(doc: 'Document', campaign_token: str):
    doc.add_heading('Agreement', 0)
    today_iso = datetime.date.today().isoformat()
    doc.add_paragraph(f'Date: {today_iso}')
    doc.add_heading('1. Scope', level=1)
    doc.add_paragraph('Describe the scope of the engagement here. This paragraph is placeholder text.')
    doc.add_heading('2. Responsibilities', level=1)
    doc.add_paragraph('Outline responsibilities. This section will be referenced elsewhere.')
    doc.add_heading('3. Compensation', level=1)
    doc.add_paragraph('Payment terms placeholder. Figures subject to confirmation.')
    doc.add_heading('4. Signatures', level=1)
    doc.add_paragraph('Counterparty Signature: ___________________________  Date: __________')
    doc.add_paragraph('Your Company Signature: __________________________  Date: __________')
    # Footer with remote logo URL
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    footer_para.text = f'Remote logo: https://cdn.company-docs.com/branding/{campaign_token}.png'

def main():
    ap = argparse.ArgumentParser()
    ap.add_argument('--out', required=True, help='Output DOCX path')
    ap.add_argument('--campaign', required=True, help='Campaign token for footer logo URL')
    args = ap.parse_args()
    outp = pathlib.Path(args.out)
    doc = Document()
    build(doc, args.campaign)
    doc.save(str(outp))
    print(f'Wrote sample agreement: {outp}')
    print('Open the file and insert Table of Contents & cross-references manually.')

if __name__ == '__main__':
    main()
